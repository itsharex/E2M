from api.core.converters.base_converter import (
    BaseConverter,
    MdElement,
    Header1,
    Header2,
    Header3,
    Paragraph,
    merge_elements_to_md,
)

from typing import List

from api.config import Config

# logging
import logging

logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)
handler = logging.StreamHandler()
handler.setLevel(logging.DEBUG)
formatter = logging.Formatter("%(asctime)s - %(name)s - %(levelname)s - %(message)s")
handler.setFormatter(formatter)
logger.addHandler(handler)


class DocxConverter(BaseConverter):

    def convert(self) -> str:
        if not self.file:
            raise ValueError("File not set")
        # todo: a better to convert docx to markdown
        elements: List[MdElement] = []

        if Config.CONVERTER == "default":
            logger.info(f"Converting [{self.file}] with default converter")
            from docx import Document

            doc = Document(self.file)

            """
            {
                "font size": {
                    "type": "header1",
                    "freq: 0.31,
                    "count": 100,
                    "paragraph": [1, 2, 3]
                },
                ...
            }

            """

            font_size_count = {}
            total_runs = 0

            for i, paragraph in enumerate(doc.paragraphs):
                for run in paragraph.runs:
                    total_runs += 1
                    font_size = run.font.size
                    if font_size not in font_size_count:
                        font_size_count[font_size] = {"count": 1, "paragraph": [i]}
                    else:
                        font_size_count[font_size]["count"] += 1
                        font_size_count[font_size]["paragraph"].append(i)

            all_font_size = set(font_size_count.keys())
            # sort the font size in descending order
            all_font_size = sorted(all_font_size, reverse=True)

            header1_font_size = 0
            header2_font_size = 0
            header3_font_size = 0

            # calculate the frequency of each font size
            for font_size in font_size_count:
                font_size_count[font_size]["freq"] = (
                    font_size_count[font_size]["count"] / total_runs
                )
                # if the frequency of a font size is greater than 0.8, it is a paragraph
                if font_size_count[font_size]["freq"] > 0.8:
                    font_size_count[font_size]["type"] = "paragraph"
                else:
                    # if the frequency of a font size is less than 0.8, it is a header
                    if len(all_font_size) >= 2 and font_size == all_font_size[0]:
                        font_size_count[font_size]["type"] = "header1"
                        header1_font_size = font_size
                    elif len(all_font_size) >= 3 and font_size == all_font_size[1]:
                        font_size_count[font_size]["type"] = "header2"
                        header2_font_size = font_size
                    elif len(all_font_size) >= 4 and font_size == all_font_size[2]:
                        font_size_count[font_size]["type"] = "header3"
                        header3_font_size = font_size

            logger.info(f"font_size_count: {font_size_count}")

            for para in doc.paragraphs:
                if not para.text.strip():
                    continue

                if para.runs[0].font.size == header1_font_size:
                    elements.append(Header1(para.text))
                elif para.runs[0].font.size == header2_font_size:
                    elements.append(Header2(para.text))
                elif para.runs[0].font.size == header3_font_size:
                    elements.append(Header3(para.text))
                else:
                    elements.append(Paragraph(para.text))

        elif Config.CONVERTER == "unstructured":
            logger.info(f"Converting [{self.file}] with unstructured converter")
            from unstructured.partition.docx import partition_docx

            unstructured_elements = partition_docx(filename=self.file)
            for element in unstructured_elements:
                # todo: need a more accurate way to determine the category
                if element.category == "Title":
                    elements.append(Header1(element.text))
                else:
                    elements.append(Paragraph(element.text))

        else:
            raise ValueError(
                f"Invalid CONVERTER: {Config.CONVERTER}, you should set\
                      CONVERTER to 'default' or 'unstructured'"
            )

        return merge_elements_to_md(elements)
